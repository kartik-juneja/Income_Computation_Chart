import tkinter as tk
from tkinter import filedialog
import json
import os
import docx
import sys
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
def change_font_in_table(table, font_name, font_size):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.name = font_name
                    font.size = Pt(font_size)
def make_text_bold_in_cell(cell):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
def set_column_alignment_to_center(table, col_index):
    for row in table.rows:
        cell = row.cells[col_index]
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
def update_listbox(path):
    listbox.insert(tk.END, path)
def getfile():
    
    file_paths = filedialog.askopenfilenames()
    path = []
    # Check if the user selected a file or canceled the dialog
    if file_paths:
        for file_path in file_paths:
            path.append(file_path)
        
    else:
        update_listbox("File selection canceled.")
    for i in range(len(path)):
# for i in range(1):
        file =open(path[i],'r')
        # file =open("c:\\Users\\sahil\\Downloads\\195067560020623.json",'r')
        x=file.read()
        # update_listbox(type(x))
        j = json.loads(x)
        # update_listbox(type(j))
        try:

            form_name=[str(key) for key in j['ITR'].keys()]
            form_name = form_name[0]
          
            year=j['ITR'][form_name]['Form_'+form_name]['AssessmentYear']
            doc=j['ITR'][form_name]['CreationInfo']['JSONCreationDate'].split('-')
            doc = '/'.join(doc[-1::-1])
            
            first_name=""
            try:
                first_name = j['ITR'][form_name]['PersonalInfo']['AssesseeName']['FirstName']
            except:
                update_listbox('Only last name update_listboxed')
            last_name = j['ITR'][form_name]['PersonalInfo']['AssesseeName']['SurNameOrOrgName']
            Father_Name=j['ITR'][form_name]['Verification']['Declaration']['FatherName']
           
            
            Name=j['ITR'][form_name]['Verification']['Declaration']['AssesseeVerName']
            # update_listbox(first_name + ' '+ last_name)
            address_value = [str(value) for value in j['ITR'][form_name]['PersonalInfo']['Address']]
            # update_listbox(j['ITR'][form_name]['PersonalInfo']['Address']['EmailAddress'])
            mobile_number = j['ITR'][form_name]['PersonalInfo']['Address']['MobileNo']
            # update_listbox(mobile_number)
            email_address = j['ITR'][form_name]['PersonalInfo']['Address']['EmailAddress']
            # update_listbox(mobile_number)
            # update_listbox(email_address)
            address = ' '.join(address_value[0:5])

            # update_listbox(address)
            PAN = j['ITR'][form_name]['PersonalInfo']['PAN']
            # update_listbox(PAN)
            DOB=j['ITR'][form_name]['PersonalInfo']['DOB'].split('-')
           
            DOB = '/'.join(DOB[-1::-1])
           
            GrossSalary=j['ITR'][form_name]['IncomeDeductions']['GrossSalary']
            BusinessIncome=j['ITR'][form_name]['IncomeDeductions']['IncomeFromBusinessProf']
            
        
            
            GrossTotalIncome=j['ITR'][form_name]['IncomeDeductions']['GrossTotIncome']
            TotalIncome=j['ITR'][form_name]['IncomeDeductions']['TotalIncome']
            StandardDeduction=j['ITR'][form_name]['IncomeDeductions']['DeductionUs16']
            OtherIncome=j['ITR'][form_name]['IncomeDeductions']['IncomeOthSrc']
            Section80C=j['ITR'][form_name]['IncomeDeductions']['UsrDeductUndChapVIA']['Section80C']
            Section80D=j['ITR'][form_name]['IncomeDeductions']['UsrDeductUndChapVIA']['Section80D']
            Section80TTA=j['ITR'][form_name]['IncomeDeductions']['UsrDeductUndChapVIA']['Section80TTA']
            SumIncome=GrossSalary+BusinessIncome+OtherIncome-StandardDeduction-Section80C-Section80D-Section80TTA
            
            if(SumIncome%10>=5):
                SumIncome=(10-(SumIncome%10))+SumIncome
            else:
                SumIncome=SumIncome-(SumIncome%10)
            
            if(TotalIncome!=SumIncome):
                HouseIncome = j['ITR'][form_name]['IncomeDeductions']['GrossRentReceived']
                RentDeduction = j['ITR'][form_name]['IncomeDeductions']['AnnualValue30Percent']
            else:
                HouseIncome = 0
                RentDeduction = 0
            
        

        
            Name_of_business=j['ITR'][form_name]['ScheduleBP']['NatOfBus44AD'][0]['NameOfBusiness']
        
            cash_in_hand=j['ITR'][form_name]['ScheduleBP']['FinanclPartclrOfBusiness']['CashInHand']
            # update_listbox(cash_in_hand)
            Refund=j['ITR'][form_name]['Refund']['RefundDue']
            Tax_Payable=j['ITR'][form_name]['TaxComputation']['TotalTaxPayable']
            Rebate=j['ITR'][form_name]['TaxComputation']['Rebate87A']
            
            TCS=j['ITR'][form_name]['TaxPaid']['TaxesPaid']['TCS']

            TDS=j['ITR'][form_name]['TaxPaid']['TaxesPaid']['TDS']



            document = docx.Document("C:\\Users\\sahil\\kartikjsonproject\\SAMPLE.docx")

            font_name = 'Courier New'
            font_size = 10
            profile_table=document.tables[0]
            profile_table.cell(0,1).text = Name
            
            make_text_bold_in_cell(profile_table.cell(0,1))
            profile_table.cell(1,1).text = Father_Name
            make_text_bold_in_cell(profile_table.cell(1,1))
            profile_table.cell(3,1).text = PAN
            make_text_bold_in_cell(profile_table.cell(3,1))
            profile_table.cell(2,3).text = str(year) + "-" + str(int(year)+1)
            profile_table.cell(3,3).text = "31/03/" + year
            profile_table.cell(4,3).text = DOB
            make_text_bold_in_cell(profile_table.cell(4,3))
            profile_table.cell(6,1).text = Name_of_business
            profile_table.cell(7,1).text = email_address
            change_font_in_table(profile_table, font_name, font_size)
        



            income_table = document.tables[1]
            income_table.cell(0,2).text = str(BusinessIncome)
            make_text_bold_in_cell(income_table.cell(0,2))
            income_table.cell(3,1).text = str(BusinessIncome)
            income_table.cell(5,1).text = str(BusinessIncome)
            income_table.cell(7,1).text = str(BusinessIncome)
            income_table.cell(9,1).text = str(BusinessIncome)
            
            income_table.cell(13,2).text = str(OtherIncome)
            income_table.cell(15,2).text = str(GrossTotalIncome)
            make_text_bold_in_cell(income_table.cell(15,2))
            income_table.cell(17,1).text = str(Section80C)
            income_table.cell(18,1).text = str(Section80D)
            income_table.cell(19,1).text = str(Section80TTA)
            income_table.cell(16,2).text = str(Section80C+Section80D+Section80TTA)
            income_table.cell(21,2).text = str(TotalIncome)
            make_text_bold_in_cell(income_table.cell(21,2))
            if(HouseIncome>0):
                income_table.cell(11,0).text = "Income from House Property"
                income_table.cell(12,0).text = "Less :-30% of Annual Value"

                income_table.cell(11,2).text = str(HouseIncome)
                income_table.cell(12,1).text = str(RentDeduction)
            else:
                income_table.cell(11,2).text = str(GrossSalary)
                income_table.cell(12,1).text = str(StandardDeduction)
                income_table.cell(11,0).text = "Income from Salary"
                income_table.cell(12,0).text = "Less :-Standard Deduction"
            # else:
            #     income_table.remove(income_table.rows[11])
            #     income_table.remove(income_table.rows[12])


            set_column_alignment_to_center(income_table, 2)
            set_column_alignment_to_center(income_table, 1)
            income_table.cell(23,2).text = "Net Assessable Income of the Assesses is thus Rs. " + str(TotalIncome)
            change_font_in_table(income_table, font_name, font_size)
            

            refund_table = document.tables[2]
            refund_table.cell(0,2).text = str(Tax_Payable)
            refund_table.cell(1,2).text = str(Rebate)
            refund_table.cell(3,1).text = str(TDS)
            refund_table.cell(4,1).text = str(TCS)
            refund_table.cell(5,1).text = str(Refund)
            
            set_column_alignment_to_center(refund_table, 2)
            set_column_alignment_to_center(refund_table, 1)
            refund_table.cell(0,0).text = "Tax on total income of Rs " + str(TotalIncome) + " at normal rates "
            refund_table.cell(1,0).text = "Rebate u/s 87a"
            refund_table.cell(6,0).text = "Rs.0 is deposited by self-assessment challan"
            change_font_in_table(refund_table, font_name, font_size)

            Date_table=document.tables[3]
            Date_table.cell(0,0).text = "Date : " + str(doc)
            make_text_bold_in_cell(Date_table.cell(0,0))
            change_font_in_table(Date_table, font_name, font_size)




            document.save(directory_path + '/' + first_name +" " + last_name + '.docx')
            l4.config(text="Computations files saved in"+directory_path)
            update_listbox(PAN + " "+ Name +" "+" Successfully created ")
        except:
            exception_type, exception_value, exception_traceback = sys.exc_info()
            update_listbox(PAN + '  ' + first_name + last_name + " Some Error Occured")
            update_listbox(f"Exception name: {exception_type.__name__}")
            update_listbox(f"Line number: {exception_traceback.tb_lineno}")
   

def on_thum_click():
    root2 = tk.Toplevel()  # Use Toplevel to create a new window
    
    image = tk.PhotoImage(file="sample_thumb.png")
    l3 = tk.Label(root2, image=image)

    l3.pack()
    root2.mainloop()

root = tk.Tk()
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()
root.geometry(f"{screen_width}x{screen_height}")
def getusername():
    try:
        username = os.getlogin()
        l3.config(text=f"User name:{username}")
    except OSError:
        update_listbox("Unable to retrieve the user name.")

def create_directory_if_not_exist(directory_path):
    if not os.path.exists(directory_path):
        os.mkdir(directory_path)
        # update_listbox("Directory created:", directory_path)
    else:
        ("Directory already exists:", directory_path)
directory_path = "C:/Users/sahil/Computation_2023-2024_KARTIK"
create_directory_if_not_exist(directory_path)


f1 = tk.Frame(root, bg='#262626')
f1.pack(fill='x')

l1 = tk.Label(f1, text="Computation chart using JSON file", font=("Helvetica", 16, "bold"))
l1.pack(pady=14)
l2 = tk.Label(root, text="Template of Computation Chart\n(Click to zoom)",font=("Helvetica", 13, "bold"))
l2.pack(pady=10)
image = tk.PhotoImage(file="sample_thumb1.png")
b3 = tk.Button(root, text="See Sample", image=image, command=on_thum_click)
b3.pack(pady=25)



b1 = tk.Button(root, text="Select file", font=("Helvetica", 12, "bold"), command=getfile)
b1.pack()

l3= tk.Label(root,text='', font=("Helvetica", 12, "bold"))
l3.pack(pady=10)

getusername()
l4= tk.Label(root,text='', font=("Helvetica", 12, "bold"))
l4.pack(pady=10)

frame = tk.Frame(root)
frame.pack(padx=20, pady=10)

listbox = tk.Listbox(root, width=500, height=100, bg='grey', fg='white', font=("Helvetica", 15, "bold"))
listbox.pack(side=tk.LEFT, fill=tk.X, expand=True)

scrollbar = tk.Scrollbar(root, orient=tk.VERTICAL, command=listbox.yview)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
listbox.config(yscrollcommand=scrollbar.set)


root.mainloop()
